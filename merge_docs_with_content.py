#!/usr/bin/env python3
"""
合并 Word 文档（处理 MHT/altChunk 格式和普通 Word 格式）
修复版 v3: 使用更健壮的 HTML 提取
"""
import sys
import os
import glob
import zipfile
import html
import quopri
import re
from xml.etree import ElementTree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Word 命名空间
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
}

def clean_html_text(text):
    """清理 HTML 文本"""
    # 移除标签
    text = re.sub(r'\u003c[^\u003e]+\u003e', ' ', text)
    # 解码实体
    text = html.unescape(text)
    # 合并多个空格
    text = re.sub(r'\s+', ' ', text)
    # 移除首尾空格
    text = text.strip()
    return text

def extract_text_from_mht(mht_content):
    """从 MHT 内容中提取文本"""
    try:
        decoded = quopri.decodestring(mht_content.encode()).decode('utf-8', errors='ignore')
        decoded = decoded.replace('=\n', '')
    except:
        decoded = mht_content
    
    texts = []
    
    # 提取标题（多种模式）
    title = ""
    title_match = re.search(r'activity-name[^\u003e]*\u003e([^\u003c]+)', decoded)
    if title_match:
        title = clean_html_text(title_match.group(1))
    if not title:
        title_match = re.search(r'rich_media_title[^\u003e]*\u003e([^\u003c]+)', decoded)
        if title_match:
            title = clean_html_text(title_match.group(1))
    if title:
        texts.append(title)
    
    # 提取正文 - 方法1: 从 js_content 或 rich_media_content 提取
    content = ""
    
    # 尝试找到 js_content 并提取其所有文本
    js_content_match = re.search(r'id=[\"\']js_content[\"\'][^\u003e]*\u003e(.*?)\u003cdiv[^\u003e]*id=', decoded, re.DOTALL)
    if js_content_match:
        content = js_content_match.group(1)
    
    # 如果失败，尝试从 js_article_content 提取
    if not content:
        js_article_match = re.search(r'id=[\"\']js_article_content[\"\'][^\u003e]*\u003e(.*)', decoded, re.DOTALL)
        if js_article_match:
            content = js_article_match.group(1)
    
    # 提取所有文本段落
    if content:
        # 提取所有 <p> 标签内容
        paragraphs = re.findall(r'\u003cp[^\u003e]*\u003e(.*?)\u003c/p\u003e', content, re.DOTALL)
        for para in paragraphs:
            text = clean_html_text(para)
            if text and len(text) > 5:
                texts.append(text)
    
    # 备用方案：如果没有提取到内容，直接提取所有文本节点
    if len(texts) <= 1:
        all_texts = re.findall(r'\u003e([^<\u003e]{20,})\u003c', decoded)
        for text in all_texts:
            text = clean_html_text(text)
            if text and len(text) > 10 and text not in texts:
                texts.append(text)
    
    return '\n\n'.join(texts)

def extract_text_from_xml(xml_content):
    """从 Word document.xml 中提取文本"""
    texts = []
    try:
        root = ET.fromstring(xml_content)
        for para in root.findall('.//w:p', NAMESPACES):
            para_texts = []
            for t in para.findall('.//w:t', NAMESPACES):
                if t.text:
                    para_texts.append(t.text)
            if para_texts:
                texts.append(''.join(para_texts))
    except Exception as e:
        print(f"XML解析错误: {e}")
    
    return '\n\n'.join(texts)

def extract_text_fallback(file_path):
    """使用 zipfile 直接读取（备用方案）"""
    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            if 'word/afchunk.mht' in z.namelist():
                mht_content = z.read('word/afchunk.mht').decode('utf-8', errors='ignore')
                return extract_text_from_mht(mht_content)
            
            if 'word/document.xml' in z.namelist():
                xml_content = z.read('word/document.xml')
                return extract_text_from_xml(xml_content)
    except Exception as e:
        print(f"备用读取失败: {e}")
    
    return ""

def extract_text_from_docx(file_path):
    """从 docx 中提取文本（支持所有格式）"""
    try:
        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text.strip())
        
        if texts and len(texts) > 0:
            return '\n\n'.join(texts)
    except Exception as e:
        pass
    
    return extract_text_fallback(file_path)

def merge_docs_with_content(folder_path, output_path):
    """合并文档，包含完整内容"""
    
    if not os.path.isdir(folder_path):
        print(f"❌ 错误: 文件夹不存在: {folder_path}")
        return False
    
    search_pattern = os.path.join(folder_path, "*.docx")
    files = sorted(glob.glob(search_pattern))
    
    if not files:
        print(f"❌ 错误: 未找到 docx 文件: {search_pattern}")
        return False
    
    print(f"📁 文件夹: {folder_path}")
    print(f"📊 找到 {len(files)} 个文档")
    print("📝 开始合并...\n")
    
    merged_doc = Document()
    success_count = 0
    empty_count = 0
    
    for i, file_path in enumerate(files, 1):
        try:
            file_name = os.path.splitext(os.path.basename(file_path))[0]
            
            print(f"  [{i}/{len(files)}] 处理: {file_name[:50]}...", end=" ")
            
            if i > 1:
                merged_doc.add_page_break()
            
            title_para = merged_doc.add_paragraph()
            title_run = title_para.add_run(file_name)
            title_run.font.size = Pt(16)
            title_run.font.bold = True
            title_run.font.color.rgb = RGBColor(0x8B, 0x5A, 0x2B)
            title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            merged_doc.add_paragraph()
            
            content = extract_text_from_docx(file_path)
            if content:
                paragraphs = content.split('\n\n')
                for para_text in paragraphs:
                    if para_text.strip():
                        para = merged_doc.add_paragraph()
                        para.add_run(para_text.strip())
                print(f"✅ ({len(paragraphs)} 段)")
                success_count += 1
            else:
                print("⚠️ 无内容")
                empty_count += 1
            
        except Exception as e:
            print(f"\n  ❌ 错误: {e}")
            continue
    
    merged_doc.save(output_path)
    print(f"\n{'='*50}")
    print(f"✅ 合并完成！")
    print(f"📄 输出文件: {output_path}")
    print(f"📊 总文档数: {len(files)}")
    print(f"✅ 成功提取: {success_count}")
    print(f"⚠️  无内容: {empty_count}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("合并 Word 文档工具（修复版 v3）")
        print("")
        print("用法:")
        print("  python3 merge_docs_with_content.py <输入文件夹> <输出文件>")
        print("")
        print("示例:")
        print("  python3 merge_docs_with_content.py ./my_docs ./merged.docx")
        sys.exit(1)
    
    folder = sys.argv[1]
    output = sys.argv[2]
    merge_docs_with_content(folder, output)
